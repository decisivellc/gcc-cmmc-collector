"""Structural analysis of policy documents (LLM Phase 0).

Extracts text from .docx / .pdf policy documents and checks for the
presence of sections expected in a CMMC-aligned policy: Purpose, Scope,
Roles and Responsibilities, Policy Statement, and Review Cycle.

This is a pre-LLM heuristic. It catches the obvious "file exists but is
a stub" failure mode. A later ticket will replace or augment this with
actual content review via a local LLM.
"""

from __future__ import annotations

import io
import logging
import re
from typing import Any

logger = logging.getLogger(__name__)


REQUIRED_SECTIONS: dict[str, list[str]] = {
    "purpose": ["purpose", "objective", "overview"],
    "scope": ["scope", "applicability", "applies to"],
    "roles": [
        "roles and responsibilities",
        "roles & responsibilities",
        "responsibilities",
        "responsibility",
        "accountable parties",
        "accountability",
        "ownership",
        "is responsible",
        "are responsible",
    ],
    "policy": [
        "policy",
        "policy statement",
        "requirements",
        "controls",
        "standards",
    ],
    "review": [
        "review",
        "review cycle",
        "review and revision",
        "annual review",
        "policy maintenance",
        "revision history",
    ],
}

MIN_WORD_COUNT_SUBSTANTIVE = 300


def extract_text(filename: str, raw: bytes) -> tuple[str | None, str | None]:
    """Return (text, error). One of them will be None."""
    if not raw:
        return None, "empty file"
    name = filename.lower()
    try:
        if name.endswith(".docx"):
            return _extract_docx(raw), None
        if name.endswith(".pdf"):
            return _extract_pdf(raw), None
        if name.endswith(".txt") or name.endswith(".md"):
            return raw.decode("utf-8", errors="replace"), None
    except Exception as exc:
        logger.warning("text extraction failed for %s: %s", filename, exc)
        return None, f"extraction error: {exc}"
    return None, "unsupported file type"


def _extract_docx(raw: bytes) -> str:
    from docx import Document  # local import so tests can mock absence
    doc = Document(io.BytesIO(raw))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _extract_pdf(raw: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(raw))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:
            logger.debug("pdf page extract error: %s", exc)
    return "\n".join(parts)


def score_document(filename: str, text: str | None, error: str | None) -> dict[str, Any]:
    """Score a document's structure. Returns a dict ready to embed in evidence."""
    if error:
        return {
            "extracted": False,
            "error": error,
            "score": 0,
            "sectionsFound": [],
            "sectionsMissing": list(REQUIRED_SECTIONS.keys()),
            "wordCount": 0,
            "substantive": False,
        }
    words = re.findall(r"\w+", text or "")
    word_count = len(words)
    normalized = _normalize_for_section_match(text or "")
    found: list[str] = []
    missing: list[str] = []
    for section, phrases in REQUIRED_SECTIONS.items():
        if any(_phrase_found(normalized, phrase) for phrase in phrases):
            found.append(section)
        else:
            missing.append(section)
    snippets: dict[str, str] = {}
    for section in missing:
        snippet = _first_snippet_containing(text or "", REQUIRED_SECTIONS[section])
        if snippet:
            snippets[section] = snippet
    return {
        "extracted": True,
        "error": None,
        "score": len(found),
        "sectionsFound": found,
        "sectionsMissing": missing,
        "wordCount": word_count,
        "substantive": word_count >= MIN_WORD_COUNT_SUBSTANTIVE,
        "missingSectionSnippets": snippets,
    }


def _first_snippet_containing(text: str, phrases: list[str]) -> str:
    """If any loose variant of the phrases appears in the raw text,
    return a 120-character snippet around the first hit. Used for
    diagnosing false-negatives: if the word IS in the text but we
    still marked it missing, the keyword list needs to grow."""
    lowered = text.casefold()
    for phrase in phrases:
        needle = phrase.casefold()
        idx = lowered.find(needle)
        if idx == -1:
            # try a shorter root for multi-word phrases
            root = needle.split()[0]
            if len(root) >= 5:
                idx = lowered.find(root)
                needle = root
        if idx != -1:
            start = max(0, idx - 40)
            end = min(len(text), idx + len(needle) + 80)
            snippet = text[start:end].replace("\n", " ").strip()
            return f"…{snippet}…"
    return ""


def _normalize_for_section_match(text: str) -> str:
    """Replace non-alphanumeric runs with single spaces, lowercase, pad with newlines."""
    lowered = text.casefold()
    cleaned = re.sub(r"[^a-z0-9\n]+", " ", lowered)
    return "\n" + cleaned + "\n"


# Section phrases should appear at a line boundary (heading-like) rather than
# buried inside a sentence. We require the phrase to be bounded by either the
# start of a line / non-word chars on the left and a newline / colon / digit
# on the right — catches "Purpose\n", "1. Purpose", "Purpose:", "Purpose 2.1".
def _phrase_found(normalized: str, phrase: str) -> bool:
    phrase_norm = re.sub(r"[^a-z0-9]+", " ", phrase.casefold()).strip()
    # Prefer heading-form matches: line starts (with optional leading numbering)
    # then the phrase then a newline / colon / end.
    heading_pattern = re.compile(
        r"(?:^|\n)\s*(?:\d+(?:\s*\d+)*\s+)?" + re.escape(phrase_norm) + r"(?:\s*\n|\s*:|\s*\d|\s*$)",
        re.MULTILINE,
    )
    if heading_pattern.search(normalized):
        return True
    # Fallback: phrase appears as a bounded token sequence anywhere (catches
    # "The purpose of this policy is ..." — still weak evidence, worth half
    # credit — but we return True and let low word count filter out stubs).
    bounded_pattern = re.compile(
        r"(?<![a-z0-9])" + re.escape(phrase_norm) + r"(?![a-z0-9])"
    )
    return bool(bounded_pattern.search(normalized))
